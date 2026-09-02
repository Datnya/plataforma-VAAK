from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

import pdfplumber
import pypdfium2 as pdfium
from docx import Document


ROOT = Path(r"C:\Users\HP\OneDrive\Desktop\Plataforma VAAK")
DOCX_DIR = ROOT / "Documentos de plataforma antigua"
WORK_DIR = ROOT / "docs" / "research" / "VAAK-RESEARCH-0-A" / "_working"
RENDERED_PDF_DIR = WORK_DIR / "rendered-pdf"
TEXT_DIR = WORK_DIR / "notebook-sources"
PAGE_IMAGE_DIR = WORK_DIR / "page-images"
SUMMARY_PATH = WORK_DIR / "corpus-summary.json"


SECRET_PATTERNS = {
    "private_key": re.compile(r"-----BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY-----"),
    "google_api_key": re.compile(r"AIza[0-9A-Za-z_-]{35}"),
    "aws_access_key": re.compile(r"\bAKIA[0-9A-Z]{16}\b"),
    "github_token": re.compile(r"\bgh[pousr]_[A-Za-z0-9]{30,}\b"),
    "bearer_token": re.compile(r"\bBearer\s+[A-Za-z0-9._~+/=-]{20,}", re.I),
    "credential_assignment": re.compile(
        r"\b(?:password|passwd|pwd|secret|token|api[_ -]?key)\s*[:=]\s*[^\s,;]{8,}", re.I
    ),
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def safe_slug(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "-", name).strip("-")


def extract_pdf(pdf_path: Path) -> tuple[list[str], list[dict]]:
    pages: list[str] = []
    page_meta: list[dict] = []
    with pdfplumber.open(pdf_path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=3, layout=True) or ""
            text = text.replace("\x00", "").strip()
            pages.append(text)
            page_meta.append(
                {
                    "page": index,
                    "characters": len(text),
                    "images": len(page.images),
                    "tables_detected": len(page.find_tables()),
                    "needs_visual_review": len(text) < 40 or bool(page.images),
                }
            )
    return pages, page_meta


def render_pdf(pdf_path: Path, output_dir: Path) -> int:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(pdf_path))
    try:
        for index in range(len(pdf)):
            page = pdf[index]
            bitmap = page.render(scale=1.5)
            bitmap.to_pil().save(output_dir / f"page-{index + 1:03d}.png")
            page.close()
        return len(pdf)
    finally:
        pdf.close()


def docx_structure(path: Path) -> dict:
    document = Document(path)
    inline_shapes = len(document.inline_shapes)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return {
        "paragraphs": len(paragraphs),
        "tables": len(document.tables),
        "inline_shapes": inline_shapes,
        "sections": len(document.sections),
        "core_title": document.core_properties.title or "",
        "core_subject": document.core_properties.subject or "",
    }


def scan_secrets(text: str) -> list[dict]:
    findings: list[dict] = []
    for label, pattern in SECRET_PATTERNS.items():
        for match in pattern.finditer(text):
            findings.append(
                {
                    "type": label,
                    "position": match.start(),
                    "preview": match.group(0)[:12] + "…",
                }
            )
    return findings


def write_notebook_source(
    logical_id: str,
    title: str,
    original: Path,
    original_hash: str,
    pages: list[str],
    page_meta: list[dict],
) -> Path:
    target = TEXT_DIR / f"{logical_id}-{safe_slug(original.stem)}.txt"
    lines = [
        f"DOCUMENTO_LOGICO: {logical_id}",
        f"TITULO: {title}",
        f"ARCHIVO_ORIGINAL: {original}",
        f"SHA256_ORIGINAL: {original_hash}",
        "METODO_EXTRACCION: PDF paginado; DOCX convertido mediante Microsoft Word cuando aplica; texto extraido con pdfplumber.",
        "NOTA_TRAZABILIDAD: Los marcadores PAGINA se refieren a la copia PDF de trabajo. Las imagenes, capturas y tablas se verifican visualmente contra las paginas renderizadas.",
        "",
    ]
    for page_number, text in enumerate(pages, start=1):
        meta = page_meta[page_number - 1]
        clean_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        cleaned = "\n".join(clean_lines)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
        lines.extend(
            [
                f"===== PAGINA {page_number} =====",
                f"[META pagina={page_number} imagenes={meta['images']} tablas_detectadas={meta['tables_detected']} revision_visual={str(meta['needs_visual_review']).lower()}]",
                cleaned if cleaned else "[SIN TEXTO EXTRAIBLE; REQUIERE REVISION VISUAL/OCR]",
                "",
            ]
        )
    target.write_text("\n".join(lines), encoding="utf-8")
    return target


def main() -> None:
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    PAGE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    sources: list[tuple[str, Path, Path]] = []
    for index, docx_path in enumerate(sorted(DOCX_DIR.glob("*.docx")), start=1):
        rendered_pdf = RENDERED_PDF_DIR / f"{docx_path.stem}.pdf"
        sources.append((f"LEG-{index:02d}", docx_path, rendered_pdf))
    sources.append(("NEW-01", ROOT / "OS&E_Requerimiento Sistema (1).pdf", ROOT / "OS&E_Requerimiento Sistema (1).pdf"))

    summary = {"ref": "VAAK-RESEARCH-0-A", "documents": [], "blocking_secret_findings": []}
    for logical_id, original, analysis_pdf in sources:
        if not analysis_pdf.exists():
            raise FileNotFoundError(f"No existe copia PDF para {original}")
        original_hash = sha256(original)
        pages, page_meta = extract_pdf(analysis_pdf)
        image_count = render_pdf(analysis_pdf, PAGE_IMAGE_DIR / logical_id)
        notebook_source = write_notebook_source(
            logical_id,
            original.stem,
            original,
            original_hash,
            pages,
            page_meta,
        )
        full_text = "\n".join(pages)
        secret_findings = scan_secrets(full_text)
        record = {
            "logical_id": logical_id,
            "original": str(original),
            "analysis_pdf": str(analysis_pdf),
            "sha256": original_hash,
            "bytes": original.stat().st_size,
            "pages": len(pages),
            "rendered_page_images": image_count,
            "characters": len(full_text),
            "pages_needing_visual_review": [m["page"] for m in page_meta if m["needs_visual_review"]],
            "page_metadata": page_meta,
            "notebook_text": str(notebook_source),
            "secret_findings": secret_findings,
        }
        if original.suffix.lower() == ".docx":
            record["docx_structure"] = docx_structure(original)
        summary["documents"].append(record)
        if secret_findings:
            summary["blocking_secret_findings"].append(
                {"logical_id": logical_id, "findings": secret_findings}
            )

    SUMMARY_PATH.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "documents": len(summary["documents"]),
        "pages": sum(d["pages"] for d in summary["documents"]),
        "characters": sum(d["characters"] for d in summary["documents"]),
        "secret_documents": len(summary["blocking_secret_findings"]),
        "summary": str(SUMMARY_PATH),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
